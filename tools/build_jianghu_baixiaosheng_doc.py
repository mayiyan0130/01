from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_TITLE = "江湖百晓生游戏框架文档"
DOC_SUBTITLE = "手机竖屏叙事经营游戏 - 世界观、系统与MVP框架 v0.1"
OUTPUT_PATH = r"C:\01\docs\江湖百晓生_游戏框架_v0.1.docx"


COLORS = {
    "ink": RGBColor(29, 45, 68),
    "gold": RGBColor(169, 123, 41),
    "rose": RGBColor(125, 49, 72),
    "mist": RGBColor(240, 244, 248),
    "sand": RGBColor(247, 242, 234),
    "line": RGBColor(214, 220, 228),
    "text": RGBColor(51, 51, 51),
    "muted": RGBColor(96, 103, 112),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_run_font(run, name="Microsoft YaHei", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def make_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.styles["Title"]
    title.font.name = "Microsoft YaHei"
    title.font.size = Pt(24)
    title.font.bold = True
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Microsoft YaHei"
    subtitle.font.size = Pt(12)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for style_name, size, color in [
        ("Heading 1", 16, COLORS["ink"]),
        ("Heading 2", 12.5, COLORS["rose"]),
        ("Heading 3", 11.5, COLORS["gold"]),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    add_text(paragraph, "江湖百晓生游戏框架文档  |  供剧情、系统与技术方案继续细化使用", size=8.5, color=COLORS["muted"])


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=24, after=6, line=1.0)
    add_text(p, DOC_TITLE, size=24, bold=True, color=COLORS["ink"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=14, line=1.0)
    add_text(p, DOC_SUBTITLE, size=12, color=COLORS["gold"])

    summary = doc.add_table(rows=3, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    widths = [Cm(3.6), Cm(10.8)]
    cover_rows = [
        ("文档用途", "固定产品骨架，作为后续剧情文案、数值表、AI接口和原型开发的共同底稿。"),
        ("产品关键词", "武侠、情报博弈、身份隐藏、AI互动、酒楼经营、手机竖屏、可攻略角色。"),
        ("当前阶段", "框架稿 v0.1。已固化核心方向；未明确细项单独标注为待确认。"),
    ]
    for row, content in zip(summary.rows, cover_rows):
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "1D2D44")
        set_cell_shading(row.cells[1], "F7F2EA")
        set_cell_margins(row.cells[0], top=100, bottom=100)
        set_cell_margins(row.cells[1], top=100, bottom=100)
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p0, after=0, line=1.1)
        set_paragraph_spacing(p1, after=0, line=1.25)
        add_text(p0, content[0], size=10.5, bold=True, color=RGBColor(255, 255, 255))
        add_text(p1, content[1], size=10.2, color=COLORS["text"])

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.autofit = False
    note.cell(0, 0).width = Cm(14.4)
    set_cell_shading(note.cell(0, 0), "F0F4F8")
    set_cell_margins(note.cell(0, 0), top=120, bottom=120, start=140, end=140)
    p = note.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.25)
    add_text(
        p,
        "核心原则: 玩家体验不是单纯读剧情，而是用“经营 + 伪装 + 情报交易 + AI角色互动”推动整个江湖运转。",
        size=11,
        bold=True,
        color=COLORS["rose"],
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(meta, before=16, after=0, line=1.0)
    add_text(meta, f"生成日期: {date.today().isoformat()}", size=9.5, color=COLORS["muted"])

    doc.add_page_break()


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(paragraph, before=4, after=8, line=1.1)
    add_text(paragraph, text, size={1: 16, 2: 12.5, 3: 11.5}[level], bold=True, color={1: COLORS["ink"], 2: COLORS["rose"], 3: COLORS["gold"]}[level])
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=6, line=1.35)
    if bold_prefix:
        add_text(paragraph, bold_prefix, size=10.5, bold=True, color=COLORS["ink"])
    add_text(paragraph, text, size=10.5, color=COLORS["text"])
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(paragraph, after=2, line=1.25)
        add_text(paragraph, item, size=10.3, color=COLORS["text"])


def add_table(doc, headers, rows, widths, header_fill="1D2D44", body_fill="FFFFFF", font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    for idx, (cell, text, width) in enumerate(zip(header.cells, headers, widths)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, top=90, bottom=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.05)
        add_text(p, text, size=9.2, bold=True, color=RGBColor(255, 255, 255))
    for row_data in rows:
        row = table.add_row()
        for cell, text, width in zip(row.cells, row_data, widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, body_fill)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            add_text(p, text, size=font_size, color=COLORS["text"])
    return table


def add_callout(doc, title, body, fill="F0F4F8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.cell(0, 0).width = Cm(16.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4, line=1.2)
    add_text(p, title, size=10.8, bold=True, color=COLORS["rose"])
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.28)
    add_text(p2, body, size=10.2, color=COLORS["text"])


def build_document():
    doc = Document()
    make_styles(doc)
    configure_page(doc)
    add_footer(doc)
    add_cover(doc)

    add_heading(doc, "1. 产品定位与设计目标", level=1)
    add_body(doc, "《江湖百晓生》定位为手机竖屏单手可玩的武侠叙事经营游戏。玩家明面上是听风阁酒楼掌柜，暗线身份则是江湖顶级情报专家“百晓生”。产品的核心卖点不是单一恋爱或单一经营，而是让“情报”成为剧情推进、身份博弈、角色攻略和酒楼建设的共同驱动力。")
    add_bullets(
        doc,
        [
            "体验支柱 1: 活的江湖。客人、门派、朝廷和随机事件持续产生新情报与新关系。",
            "体验支柱 2: 身份张力。玩家并非公开无敌，暴露风险会真实影响玩法节奏。",
            "体验支柱 3: 攻略可经营。角色关系不是独立恋爱线，而是能反哺情报、经营和主线。",
            "体验支柱 4: 轻操作插入。三消、二合、贪吃蛇不是割裂小游戏，而是剧情和系统的结算接口。",
        ],
    )
    add_callout(doc, "一句话定位", "玩家以酒楼掌柜的身份潜伏江湖中心，在武林大比前后的乱局中经营听风阁、交易情报、周旋于正邪朝堂之间，并借由AI角色互动塑造自己的江湖声名。")

    add_heading(doc, "2. 世界观、舞台与势力结构", level=1)
    add_body(doc, "当前故事大背景是“四年一度武林大比”将至，各门派、朝廷与魔教势力同时入局。武林大比不仅决定新一任武林盟主，也会影响各派调令、资源分配和朝堂对江湖的控制方式。拜月与五毒首次参赛，使局面更具火药味。")
    add_table(
        doc,
        headers=["势力", "立场", "对听风阁的价值", "当前敏感点"],
        rows=[
            ("朝廷", "官方势力", "提供法令、税务、调令和政治情报来源。", "对武林大比高度关注，可能暗中操盘。"),
            ("武当", "正派", "提供正统江湖风向与礼法背书。", "与朝廷合作尺度需观察。"),
            ("少林", "正派", "提供稳重、权威、调停型信息。", "可能在各派冲突中扮演压舱石。"),
            ("华山", "正派", "盟主与旧同门线高度集中。", "江潋与玩家旧事将直接影响主线张力。"),
            ("峨嵋", "正派", "女侠线、门派振兴线、正道声望线关键。", "谢扶摇对责任的执念会带来任务压力。"),
            ("唐门", "中立偏复杂", "暗器、机关、密探类情报价值高。", "擅长隐藏，适合设计反情报对抗。"),
            ("药王谷", "中立", "伤药、毒理、疗养与稀有物资线。", "可与经营和救人剧情强绑定。"),
            ("拜月", "邪派/被视作魔教", "高风险高价值信息与身份试探线。", "理念冲突大，不等于角色一定为恶。"),
            ("五毒", "邪派/被视作魔教", "毒术、虫蛊、边境传闻与隐秘委托。", "与拜月联动会显著提高悬疑感。"),
        ],
        widths=[Cm(2.4), Cm(2.6), Cm(6.0), Cm(5.0)],
        body_fill="F7F9FC",
    )
    add_body(doc, "原则上，门派“正/中立/邪”只代表外界认知与组织立场，不直接等价于善恶判断。这个设定必须在角色塑造和AI输出里长期维持，避免人物脸谱化。")

    add_heading(doc, "3. 玩家身份与核心资源", level=1)
    add_body(doc, "玩家姓名可自定义，默认名为“言笑笑”；性别可选男/女。玩家在江湖上极少以真容示人，多数人只知道听风阁掌柜，不知道掌柜本人就是百晓生。这个“双层身份”是全系统的中心。")
    add_table(
        doc,
        headers=["资源", "主要来源", "主要用途", "失衡风险"],
        rows=[
            ("银两", "卖情报、酒楼经营、任务奖励", "升级酒楼、购买道具、解锁设施", "收入不足会卡经营成长"),
            ("情报条目", "对话、事件、探索、贪吃蛇收集", "整合、出售、触发大事件分支", "低质量情报会误导判断"),
            ("情报完整度", "碎片整合、关联线索补齐", "提升售价、解锁真相、开启隐藏剧情", "错过时限会降值"),
            ("暴露风险", "说错话、被试探、身份痕迹外露", "驱动危机事件与三消减压玩法", "过高会触发身份公开或强制追杀"),
            ("人脉/信任", "正确回应、完成委托、长期陪伴", "降低试探、解锁专属剧情与特权", "立场摇摆会导致关系断裂"),
            ("酒楼繁荣度", "装修升级、菜品、员工配置、口碑", "提升客流、刷新高价值客人", "经营弱会影响剧情承接密度"),
        ],
        widths=[Cm(2.5), Cm(4.2), Cm(4.8), Cm(4.5)],
        body_fill="FFFFFF",
    )

    add_heading(doc, "4. 核心循环与系统联动", level=1)
    add_body(doc, "建议把单日/单阶段循环稳定为“接客 -> 对话打探 -> 事件触发 -> 小游戏结算 -> 情报整合/售卖 -> 酒楼建设 -> 次日刷新”。这样可以保证剧情、经营和小游戏处在同一闭环内，而不是互相打断。")
    add_table(
        doc,
        headers=["阶段", "玩家动作", "主要产出", "会影响的后续系统"],
        rows=[
            ("客人刷新", "接待固定NPC与随机客人", "新关系、新线索、新委托", "关系线、随机事件池"),
            ("对话周旋", "询问、试探、安抚、隐瞒", "情报碎片、态度变化、风险变化", "AI对话、身份系统"),
            ("事件触发", "介入武林、朝堂或私人事件", "支线开启、小游戏入口", "主线推进、资源波动"),
            ("结算与整合", "完成三消/二合/贪吃蛇等结算", "风险降低、线索升级、奖励提升", "情报价值、声望、角色观感"),
            ("经营回合", "升级酒楼、配置帮工与设施", "客流与酒楼数值提升", "次日刷新质量、商店与玩法深度"),
        ],
        widths=[Cm(2.2), Cm(4.4), Cm(4.2), Cm(5.0)],
        body_fill="F7F9FC",
    )
    add_callout(
        doc,
        "大事件情报机制",
        "当遇到武林大比、门派调令、朝堂密令等大事件时，玩家会获得若干“情报碎片”。若在限定时段内把同一事件的全部关键信报整合完成，则该事件情报价值翻倍，并能额外打开隐藏结论、专属卖家或特殊剧情分支。",
        fill="F7F2EA",
    )

    add_heading(doc, "5. AI叙事接口框架", level=1)
    add_body(doc, "AI的职责应聚焦在“生成自然对话、补足情绪表达、根据状态差异变体叙事”，而不应直接掌控所有数值结算。数值、事件解锁条件、身份暴露阈值和奖励分发，必须由规则层先判断，再把结果作为上下文喂给AI。")
    add_table(
        doc,
        headers=["模块", "由规则层负责", "由AI负责"],
        rows=[
            ("对话入口", "当前事件ID、NPC身份、关系值、风险值、可说话题", "生成符合人设的台词与反应"),
            ("情绪与好感", "数值增减、标签变化、触发条件", "把变化翻译成自然的态度差异"),
            ("情报生成", "情报类别、真假比例、来源可信度", "按口吻包装信息并留下可追踪线索"),
            ("隐藏身份试探", "是否触发、成功失败判定、风险增减", "输出试探性问题、套话和心理压迫感"),
            ("剧情收束", "分支选择、结算奖励、存档写入", "写出符合分支结果的剧情片段"),
        ],
        widths=[Cm(2.8), Cm(6.0), Cm(6.0)],
        body_fill="FFFFFF",
    )
    add_bullets(
        doc,
        [
            "建议保留结构化状态槽: `player_public_role`、`hidden_identity_exposure`、`npc_trust`、`npc_affection`、`faction_stance`、`current_event_stage`。",
            "建议每次AI调用都带上“不可改写字段”: 玩家已知事实、NPC是否知晓真实身份、门派立场、事件时间点。",
            "建议输出采用“台词 + 情绪标签 + 是否新增情报 + 是否触发风险”的结构，方便前端和服务端消费。",
            "不建议让AI直接编造新规则、新门派关系或新身份真相；世界规则应由配置表控制。",
        ],
    )

    add_heading(doc, "6. 小游戏挂接设计", level=1)
    add_table(
        doc,
        headers=["玩法", "触发场景", "成功收益", "失败代价", "设计备注"],
        rows=[
            ("三消", "说错话、被怀疑、真实身份有暴露风险时", "显著降低暴露风险，并保住当前对话局势", "风险继续累积，可能进入更高压事件", "用于表现临场圆谎、转移注意和压制流言"),
            ("贪吃蛇", "潜入打探、街巷收集、客栈外情报追逐事件", "规定时间内吃到越多，得到的情报碎片越多", "只能获得部分线索，真相链条会断", "适合表达“追索”和“收网”感"),
            ("二合", "建议用于情报归档升级、酒楼物资合成或道具强化 (待确认)", "提升情报等级、设施品质或稀有道具阶级", "浪费材料或延迟升级效率", "该玩法定位建议尽快拍板，否则系统边界会持续模糊"),
        ],
        widths=[Cm(1.8), Cm(4.5), Cm(3.8), Cm(3.8), Cm(3.4)],
        body_fill="F7F9FC",
        font_size=9.0,
    )
    add_body(doc, "当前最清晰的挂接关系已经成立: 三消服务于“身份危机缓释”；贪吃蛇服务于“情报收集效率”；二合则建议服务于“经营升级或情报升阶”。如果二合继续悬空，后续数值、经济和UI框架都会被拖慢。")

    add_heading(doc, "7. 固定NPC与攻略逻辑", level=1)
    add_table(
        doc,
        headers=["NPC", "身份", "核心攻略属性", "是否知晓真实身份", "主要收益", "风险与剧情钩子"],
        rows=[
            ("莫红绫", "帮工 / 前拜月左护法“赤月狐”", "羁绊 / 忠诚 / 陪伴", "知晓", "玩家最稳定的后方支援与护卫线", "若玩家暴露，她会成为最后防线；醋意和嘴硬可形成高频互动"),
            ("南宫翊", "当朝三皇子", "崇拜 / 资源 / 直球", "不知晓", "可解锁朝廷资源、税收减免、特权渠道", "对百晓生有幻想，揭穿真相时反差会很强"),
            ("谢扶摇", "峨嵋大师姐“清音仙子”", "信任 / 守护 / 正直", "不知晓", "可获得正派情报与峨嵋关系加成", "会在责任与感情之间拉扯；过度逞强是主要事件点"),
            ("江潋", "武林盟主 / 华山掌门", "占有 / 宿命 / 旧情", "大概率可逐步察觉", "主线情感张力最强，能牵动华山与盟主线", "若玩家为女性且互动暧昧，可触发“华山旧事”隐藏剧情"),
            ("晏无秋", "拜月圣子“天狼”", "危险 / 探求 / 试探", "不知晓但高度怀疑", "可打开高风险高价值的魔教视角情报", "与其对话最容易触发暴露风险和三消危机"),
        ],
        widths=[Cm(2.0), Cm(3.3), Cm(3.3), Cm(2.5), Cm(3.1), Cm(3.7)],
        body_fill="FFFFFF",
        font_size=8.9,
    )
    add_body(doc, "除固定NPC外，还应搭建“随机人物 + 随机门派客人 + 特殊事件访客”池。AI对话最适合先服务这一层，因为随机NPC能显著放大“江湖是活的”这条卖点，而不必把所有内容制作压力压在固定角色身上。")

    add_heading(doc, "8. 主线时间轴: 武林大比赛季结构", level=1)
    add_table(
        doc,
        headers=["阶段", "剧情重心", "系统重心", "玩家目标"],
        rows=[
            ("阶段1 预热", "各派入城，听风阁开始热闹", "基础经营、客人刷新、初识NPC", "站稳酒楼，建立第一批关系"),
            ("阶段2 暗潮", "门派互探、朝廷布线、拜月与五毒入局", "情报搜集、贪吃蛇事件增多", "判断谁在撒谎、谁值得押注"),
            ("阶段3 危机", "玩家身份被试探，大事件密集爆发", "三消危机、情报整合、路线抉择", "压低暴露风险并抓住核心情报"),
            ("阶段4 争位", "武林大比进入关键回合", "高价值委托、势力站队、专属剧情", "让自己支持的结局逐步成型"),
            ("阶段5 余波", "盟主结果落地，朝堂与江湖重新洗牌", "结算经营成果、关系成果和真相线", "决定听风阁与百晓生的江湖位置"),
        ],
        widths=[Cm(2.6), Cm(4.7), Cm(4.4), Cm(4.3)],
        body_fill="F7F9FC",
    )

    add_heading(doc, "9. 技术与内容生产框架建议", level=1)
    add_table(
        doc,
        headers=["层级", "建议职责", "落地说明"],
        rows=[
            ("客户端", "竖屏UI、剧情演出、酒楼经营页、小游戏模块、存档表现", "小游戏建议做成独立子模块，统一由事件系统调用"),
            ("规则服务", "事件调度、数值判定、经济结算、身份风险判定", "所有关键结算先过规则层，避免AI污染状态"),
            ("AI编排服务", "拼接上下文、调用模型、校验输出结构、写入短期记忆", "必须保留人设卡、事实卡和禁改字段"),
            ("内容配置", "NPC卡、门派卡、事件卡、情报卡、商店与酒楼配置", "尽量表驱动，后期扩剧情和角色更稳"),
            ("数据实体", "玩家档案、关系状态、情报库、酒楼状态、赛季进度", "这些实体会决定后续接口和存档格式"),
        ],
        widths=[Cm(2.4), Cm(5.5), Cm(7.8)],
        body_fill="FFFFFF",
    )
    add_callout(
        doc,
        "建议先做MVP，不要一开始就铺满整片江湖",
        "第一版只需要跑通 1 个赛季主线、5 个固定可攻略角色、1 套酒楼经营基础循环、3 种小游戏入口、1 套身份暴露系统。只要这套闭环成立，后面追加门派、角色和事件的成本才可控。",
        fill="F7F2EA",
    )

    add_heading(doc, "10. 当前待确认事项", level=1)
    add_bullets(
        doc,
        [
            "二合玩法究竟绑定“酒楼物资升级”还是“情报归档升阶”。这个决定会直接影响经济系统结构。",
            "玩家性别分支的深度范围。当前已知江潋线存在性别触发差异，其他角色是否也要做显式分支需要尽快决定。",
            "AI接口使用范围。是只覆盖自由对话，还是连随机事件文案和部分支线演出也一起接入。",
            "随机NPC池的制作策略。建议先做模板化身份池，再让AI做口吻差异，而不是全随机生成人设。",
            "酒楼经营深度是否只做轻装修与设施升级，还是加入菜谱、员工排班、房间经营等更重系统。",
        ],
    )
    add_body(doc, "如果下一步继续推进，最合理的顺序是: 先把MVP系统框图和数据表字段定下来，再拆剧情主线章节，最后再扩写各角色攻略线。这样能避免文案先跑太远，导致系统承接不上。")

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
